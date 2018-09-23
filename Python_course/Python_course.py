#陈友文文字检索

'''
制作字典 word
    1，扫描ROI
    2，去掉分页符，分节符 回车，分栏符,变成一整段,word
    3，单个的零，大小写o转化为句号，word
    5，句号转化为分节符 word
    5，删除没有”词汇注释“的段落，，程序
    4，黑体字前加换行，word
    7，检索，有则返回字符串，没有则返回单词加换行，程序

'''

from docx import Document
from docx.shared import Inches
from fnmatch import fnmatch, fnmatchcase
#定义全局变量
word ="sphere1"#目标词
result=[]#保存位置
#读取文件
result=open("result.txt","wt")
#--------------------------清理函数，去掉无关段落 】------------------------------------------------------

def qingli():
    doc=Document('讲解.docx')
    for p in doc.paragraphs:
        if p.text.find("词汇注释")==-1:        
            p.clear()
            print (p.text)
    doc.save("词汇注释.docx")

#--------------------------清理函数，去掉无关段落 】------------------------------------------------------

def Search (word):
    str1=""
    doc=Document("词汇注释.docx")
    for word1 in [word,word[0:-2],word[0:-3],word[-4]]:
        for i in doc.paragraphs:
            if i.text.startswith(word1)!=-1:
                str1=i.text+"\n"
                break
    if len(str1)>1:
        if str1.find('记')!=-1:
            str1=str1.replace('记', "\n")
        result.write(word+'\n'+str1)
        print(word+'\n'+str1)
    else:
        str1="……没有找到"+word+"\n"
        result.write(str1)
        print(str1)
#--------------------------清理函数，去掉无关段落 】------------------------------------------------------

def startSearch (word):
    str1=""
    doc=Document("词汇注释.docx")
    for word1 in [word,word[0:-1],word[0:-2],word[0:-3]]:
        for i in doc.paragraphs:
            if fnmatch(i.text, word1+'*'):              
                str1=i.text+"\n" 
                break              
    if len(str1)>1:
        result.write(word+'\n'+str1)
        print(word+'\n'+str1)
    else:
        str1="……"+word+"\n"
        result.write(str1)
        print(str1)
#获取要查的单词列表
def getText():
    txt=open("word.txt","r").read()
    txt=txt.replace('\n', " ")
    datals=txt.split()
    return datals
def main():

    datals=getText()
    for i in datals:
        startSearch(i)
    print(datals)
   # qingli()
main()
#for lines in word:
#    datals.append(lines)
#word.close()
#print(datals)
#Search("spheres")



#复制到结束点
#保存，退出

#--------------------------【 找黑体 】------------------------------------------------------
#for p in doc.paragraphs:
#    for r in p.runs:#
#        if r.bold:         #找黑体           
#            print(r.text)



'''
for p in doc.paragraphs:
    if word in p.text :#找到每一句开始点   
        print(p.text[p.text.find(word):p.text.find(" 例 ")])
        print()
        print()
'''